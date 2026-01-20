"""
Export Module.
Handles exporting Q&A sessions to PDF and Word formats.
"""

from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime
from io import BytesIO

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import EXPORTS_DIR


class SessionExporter:
    """
    Export Q&A sessions to PDF and Word formats.
    """
    
    def __init__(self):
        """Initialize exporter."""
        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    
    def export_to_pdf(
        self,
        messages: List[Dict[str, Any]],
        filename: str = None
    ) -> Path:
        """
        Export session to PDF.
        
        Args:
            messages: List of chat messages
            filename: Optional custom filename
            
        Returns:
            Path to exported PDF
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"swiggy_qa_{timestamp}.pdf"
        
        filepath = EXPORTS_DIR / filename
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('#FC8019'),
            spaceAfter=30
        )
        
        question_style = ParagraphStyle(
            'Question',
            parent=styles['Normal'],
            fontSize=12,
            textColor=HexColor('#333333'),
            leftIndent=20,
            spaceBefore=20,
            spaceAfter=10,
            borderPadding=10
        )
        
        answer_style = ParagraphStyle(
            'Answer',
            parent=styles['Normal'],
            fontSize=11,
            textColor=HexColor('#555555'),
            leftIndent=20,
            spaceBefore=5,
            spaceAfter=15
        )
        
        citation_style = ParagraphStyle(
            'Citation',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor('#888888'),
            leftIndent=40
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("Swiggy Annual Report Q&A Session", title_style))
        story.append(Paragraph(
            f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}",
            styles['Normal']
        ))
        story.append(Spacer(1, 30))
        
        # Messages
        for msg in messages:
            if msg["role"] == "user":
                story.append(Paragraph(f"<b>Question:</b> {msg['content']}", question_style))
            else:
                story.append(Paragraph(f"<b>Answer:</b>", styles['Heading3']))
                story.append(Paragraph(msg['content'], answer_style))
                
                # Confidence
                confidence = msg.get('confidence', 0)
                story.append(Paragraph(
                    f"<i>Confidence: {int(confidence * 100)}%</i>",
                    citation_style
                ))
                
                # Citations
                citations = msg.get('citations', [])
                if citations:
                    story.append(Paragraph("<b>Sources:</b>", citation_style))
                    for c in citations:
                        story.append(Paragraph(
                            f"• Page {c['page']} - {c['section']}",
                            citation_style
                        ))
                
                story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    def export_to_word(
        self,
        messages: List[Dict[str, Any]],
        filename: str = None
    ) -> Path:
        """
        Export session to Word document.
        
        Args:
            messages: List of chat messages
            filename: Optional custom filename
            
        Returns:
            Path to exported Word document
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"swiggy_qa_{timestamp}.docx"
        
        filepath = EXPORTS_DIR / filename
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading("Swiggy Annual Report Q&A Session", 0)
        title.runs[0].font.color.rgb = RGBColor(252, 128, 25)
        
        # Timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        timestamp_para.runs[0].italic = True
        
        doc.add_paragraph()  # Spacer
        
        # Messages
        for msg in messages:
            if msg["role"] == "user":
                q_para = doc.add_paragraph()
                q_run = q_para.add_run("Question: ")
                q_run.bold = True
                q_para.add_run(msg['content'])
            else:
                # Answer heading
                a_heading = doc.add_paragraph()
                a_heading.add_run("Answer:").bold = True
                
                # Answer content
                doc.add_paragraph(msg['content'])
                
                # Confidence
                confidence = msg.get('confidence', 0)
                conf_para = doc.add_paragraph()
                conf_run = conf_para.add_run(f"Confidence: {int(confidence * 100)}%")
                conf_run.italic = True
                conf_run.font.size = Pt(10)
                
                # Citations
                citations = msg.get('citations', [])
                if citations:
                    sources_para = doc.add_paragraph()
                    sources_para.add_run("Sources:").bold = True
                    
                    for c in citations:
                        cite_para = doc.add_paragraph(style='List Bullet')
                        cite_para.add_run(f"Page {c['page']} - {c['section']}")
                        cite_para.paragraph_format.left_indent = Inches(0.5)
                
                doc.add_paragraph()  # Spacer
        
        # Save
        doc.save(str(filepath))
        
        return filepath
    
    def export_to_bytes_pdf(self, messages: List[Dict[str, Any]]) -> bytes:
        """Export to PDF and return bytes (for Streamlit download)."""
        buffer = BytesIO()
        
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = [
            Paragraph("Swiggy Annual Report Q&A Session", styles['Heading1']),
            Spacer(1, 20)
        ]
        
        for msg in messages:
            if msg["role"] == "user":
                story.append(Paragraph(f"<b>Q:</b> {msg['content']}", styles['Normal']))
            else:
                story.append(Paragraph(f"<b>A:</b> {msg['content']}", styles['Normal']))
                story.append(Spacer(1, 10))
        
        doc.build(story)
        
        buffer.seek(0)
        return buffer.getvalue()


# CLI for testing
if __name__ == "__main__":
    exporter = SessionExporter()
    
    # Test messages
    messages = [
        {"role": "user", "content": "What was Swiggy's revenue?"},
        {
            "role": "assistant",
            "content": "Swiggy's total revenue was Rs. 10,000 crores in FY 2023-24.",
            "confidence": 0.85,
            "citations": [
                {"page": 42, "section": "Financial Statements", "relevance_score": 0.9}
            ]
        }
    ]
    
    # Export
    pdf_path = exporter.export_to_pdf(messages)
    print(f"Exported PDF to: {pdf_path}")
    
    word_path = exporter.export_to_word(messages)
    print(f"Exported Word to: {word_path}")
